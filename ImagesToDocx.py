import os
from docx import Document
from docx.shared import Inches
from PIL import Image

def add_images_to_word(image_folder, output_doc):
    # Create a new Document
    doc = Document()

    # Get list of image files from the folder
    image_files = [f for f in os.listdir(image_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif'))]

    # Loop through the images and add each as a new full-page image in the Word doc
    for image_file in image_files:
        image_path = os.path.join(image_folder, image_file)
        
        # Open the image and rotate it 90 degrees clockwise
        img = Image.open(image_path)
        rotated_img = img.rotate(-90, expand=True)  # Negative angle for clockwise rotation

        # Save the rotated image temporarily
        rotated_image_path = os.path.join(image_folder, f"rotated_{image_file}")
        rotated_img.save(rotated_image_path)
        img.close()

        # Add a new section if not the first page
        if doc.paragraphs:
            doc.add_page_break()

        # Add the rotated image to the Word doc
        doc.add_picture(rotated_image_path, width=Inches(6))  # Adjust size to fit within a typical Word page (A4 width)

        # Remove the temporary rotated image after it's been added
        os.remove(rotated_image_path)

    # Save the document
    doc.save(output_doc)
    print(f"Document '{output_doc}' created with {len(image_files)} rotated images.")

# Example usage
image_folder = 'Scripture'
output_doc = 'Scripture.docx'
add_images_to_word(image_folder, output_doc)
